from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.services.rag.parsing.base import (
    DocumentParser,
    ParsedBlock,
    ParsedDocument,
    ParsedImage,
)
from app.services.rag.parsing.image_store import ImageStore
from app.services.rag.parsing.md_parser import glue_images_to_paragraphs

_EMBED_ID_RE = re.compile(r'rId\d+')
_IMAGE_REF_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_DATA_INDEX_SECTION_KEYWORD = "数据索引"
_NUMBERED_HEADING_RE = re.compile(r"^\s*\d+[\.\)、．]\s*\S")
_TITLE_STYLE_PREFIXES = ("title", "标题", "subtitle", "副标题")
_MAX_HEADING_CHARS = 120


def _paragraph_plain_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs if run.text).strip()


def _is_bold_short_heading(paragraph: Paragraph, text: str) -> bool:
    if not text or len(text) > 80:
        return False
    runs = [run for run in paragraph.runs if (run.text or "").strip()]
    if not runs:
        return False
    bold_runs = [run for run in runs if run.bold is True]
    if not bold_runs:
        return False
    bold_chars = sum(len(run.text) for run in bold_runs)
    return bold_chars >= max(len(text) * 0.6, 1)


def _is_heading_paragraph(paragraph: Paragraph) -> bool:
    style = paragraph.style
    style_name = (style.name if style is not None else None) or ""
    if style_name.startswith("Heading"):
        return True
    lowered_style = style_name.casefold()
    if any(lowered_style.startswith(prefix) for prefix in _TITLE_STYLE_PREFIXES):
        return True
    text = _paragraph_plain_text(paragraph)
    if not text or len(text) > _MAX_HEADING_CHARS:
        return False
    if _NUMBERED_HEADING_RE.match(text):
        return True
    if _is_bold_short_heading(paragraph, text):
        return True
    return False


def _heading_label(paragraph: Paragraph) -> str:
    text = _paragraph_plain_text(paragraph)
    if not text:
        return ""
    style = paragraph.style
    style_name = (style.name if style is not None else None) or ""
    if style_name.startswith("Heading") or _NUMBERED_HEADING_RE.match(text):
        return text
    if _is_bold_short_heading(paragraph, text):
        return text
    return text


def _suffix_from_content_type(content_type: str | None) -> str:
    if not content_type:
        return "bin"
    if "/" in content_type:
        return content_type.rsplit("/", 1)[-1].lower() or "bin"
    return content_type.lstrip(".").lower() or "bin"


def _embed_ids_from_element(element: object) -> list[str]:
    ids: list[str] = []
    for blip in element.xpath(".//a:blip"):  # type: ignore[attr-defined]
        embed = blip.get(qn("r:embed"))
        if embed and embed not in ids:
            ids.append(embed)
    if not ids:
        xml = getattr(element, "xml", "") or ""
        for match in _EMBED_ID_RE.findall(xml):
            if match not in ids:
                ids.append(match)
    return ids


def _image_markdown(*, alt: str, storage_key: str) -> str:
    return f"![{alt}]({storage_key})"


def _is_data_index_section_heading(heading: str) -> bool:
    return _DATA_INDEX_SECTION_KEYWORD in heading


def _normalize_paragraph_layout(text: str) -> str:
    """将 Word 中「图+制表符+说明」整理为 Markdown 友好的多行结构。"""
    if not text:
        return text
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\t", "\n")
    normalized = re.sub(r"(!\[[^\]]*\]\([^)]+\))", r"\1\n\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    normalized = glue_images_to_paragraphs(normalized.strip())
    return _dedupe_image_refs_in_text(normalized)


def _dedupe_image_refs_in_text(text: str) -> str:
    seen_keys: set[str] = set()

    def replace(match: re.Match[str]) -> str:
        key = match.group(2).strip()
        if key in seen_keys:
            return ""
        seen_keys.add(key)
        return match.group(0)

    deduped = _IMAGE_REF_RE.sub(replace, text)
    return re.sub(r"\n{3,}", "\n\n", deduped).strip()


class DocxDocumentParser(DocumentParser):
    extensions = (".docx",)

    def __init__(
        self,
        *,
        extract_images: bool = True,
        image_store: ImageStore | None = None,
    ) -> None:
        self._extract_images = extract_images
        self._image_store = image_store or ImageStore()

    def parse(self, path: str | Path) -> ParsedDocument:
        file_path = Path(path)
        try:
            document = docx.Document(str(file_path))
            return self._to_document(document)
        except Exception as exc:
            raise ValueError(
                f"无法解析 DOCX 文件 {file_path.name}: {exc}"
            ) from exc

    def _store_image(
        self,
        *,
        document: DocxDocument,
        r_id: str,
        cache: dict[str, str],
        images: list[ParsedImage],
    ) -> str | None:
        if r_id in cache:
            return cache[r_id]
        try:
            part = document.part.related_parts[r_id]
        except KeyError:
            return None
        content_type = getattr(part, "content_type", "") or ""
        if not content_type.startswith("image/"):
            return None
        blob = part.blob
        if not blob:
            return None
        storage_key = self._image_store.save(
            blob,
            suffix=_suffix_from_content_type(content_type),
        )
        cache[r_id] = storage_key
        images.append(
            ParsedImage(
                page=None,
                storage_key=storage_key,
                index_in_page=len(images),
            )
        )
        return storage_key

    def _inline_images_markdown(
        self,
        element: object,
        *,
        document: DocxDocument,
        cache: dict[str, str],
        images: list[ParsedImage],
        alt: str,
    ) -> str:
        if not self._extract_images:
            return ""
        lines: list[str] = []
        for r_id in _embed_ids_from_element(element):
            storage_key = self._store_image(
                document=document,
                r_id=r_id,
                cache=cache,
                images=images,
            )
            if storage_key is not None:
                lines.append(_image_markdown(alt=alt, storage_key=storage_key))
        return "\n".join(lines)

    def _paragraph_text(
        self,
        paragraph: Paragraph,
        *,
        document: DocxDocument,
        cache: dict[str, str],
        images: list[ParsedImage],
    ) -> str:
        parts: list[str] = []
        for run in paragraph.runs:
            text = run.text
            if text:
                parts.append(text)
            if self._extract_images:
                inline = self._inline_images_markdown(
                    run._element,
                    document=document,
                    cache=cache,
                    images=images,
                    alt="图示",
                )
                if inline:
                    parts.append(inline)
        body = "".join(parts).strip()
        if not body and self._extract_images:
            body = self._inline_images_markdown(
                paragraph._element,
                document=document,
                cache=cache,
                images=images,
                alt="图示",
            ).strip()
        return _normalize_paragraph_layout(body)

    def _table_text(
        self,
        table: Table,
        *,
        document: DocxDocument,
        cache: dict[str, str],
        images: list[ParsedImage],
    ) -> str:
        rows: list[str] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_lines: list[str] = []
                for paragraph in cell.paragraphs:
                    content = self._paragraph_text(
                        paragraph,
                        document=document,
                        cache=cache,
                        images=images,
                    )
                    if content:
                        cell_lines.append(content)
                if cell_lines:
                    cells.append("\n".join(cell_lines))
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _append_orphan_images(
        self,
        document: DocxDocument,
        *,
        cache: dict[str, str],
        images: list[ParsedImage],
        blocks: list[ParsedBlock],
    ) -> None:
        orphan_lines: list[str] = []
        for rel in document.part.rels.values():
            if "image" not in rel.reltype:
                continue
            r_id = rel.rId
            if r_id in cache:
                continue
            storage_key = self._store_image(
                document=document,
                r_id=r_id,
                cache=cache,
                images=images,
            )
            if storage_key is not None:
                orphan_lines.append(_image_markdown(alt="图示", storage_key=storage_key))
        if not orphan_lines:
            return
        appendix = glue_images_to_paragraphs("\n\n".join(orphan_lines))
        if blocks:
            last = blocks[-1]
            merged = glue_images_to_paragraphs(
                f"{last.text}\n\n{appendix}".strip()
            )
            blocks[-1] = ParsedBlock(
                page=last.page,
                text=merged,
                heading=last.heading,
            )
        else:
            blocks.append(ParsedBlock(page=None, text=appendix, heading=None))

    def _to_document(self, document: DocxDocument) -> ParsedDocument:
        blocks: list[ParsedBlock] = []
        images: list[ParsedImage] = []
        cache: dict[str, str] = {}

        current_heading: str | None = None
        current_body: list[str] = []
        skipping_data_index = False

        def flush() -> None:
            nonlocal current_heading, current_body
            text = _normalize_paragraph_layout("\n\n".join(current_body).strip())
            if not text:
                current_body = []
                return
            blocks.append(
                ParsedBlock(page=None, text=text, heading=current_heading)
            )
            current_heading = None
            current_body = []

        iter_blocks = (
            document.iter_inner_content()
            if hasattr(document, "iter_inner_content")
            else list(document.paragraphs)
        )
        for block in iter_blocks:
            if isinstance(block, Paragraph):
                content = self._paragraph_text(
                    block,
                    document=document,
                    cache=cache,
                    images=images,
                )
                if not content or (
                    len(content) <= 2 and not any(character.isalnum() for character in content)
                ):
                    continue
                if _is_heading_paragraph(block):
                    label = _heading_label(block) or content
                    if _is_data_index_section_heading(label):
                        flush()
                        skipping_data_index = True
                        current_heading = None
                        current_body = []
                        continue
                    skipping_data_index = False
                    flush()
                    current_heading = label
                    continue
                if skipping_data_index:
                    continue
                current_body.append(content)
            elif isinstance(block, Table):
                table_text = self._table_text(
                    block,
                    document=document,
                    cache=cache,
                    images=images,
                )
                if table_text:
                    current_body.append(table_text)

        flush()

        if self._extract_images:
            self._append_orphan_images(
                document,
                cache=cache,
                images=images,
                blocks=blocks,
            )

        if not blocks:
            fallback = "\n".join(
                p.text.strip() for p in document.paragraphs if p.text.strip()
            )
            if fallback:
                blocks.append(
                    ParsedBlock(page=None, text=fallback, heading=None)
                )

        full_text = glue_images_to_paragraphs(
            "\n\n".join(
                (f"## {block.heading}\n\n{block.text}" if block.heading else block.text)
                for block in blocks
                if block.text.strip()
            )
        )
        return ParsedDocument(full_text=full_text, blocks=blocks, images=images)
