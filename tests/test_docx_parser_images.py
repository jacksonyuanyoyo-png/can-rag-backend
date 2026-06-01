from __future__ import annotations

from io import BytesIO
from pathlib import Path

import docx

from app.core.config import Settings
from app.domain.import_job import ChunkingConfig
from app.services.rag.parsing.docx_parser import DocxDocumentParser
from app.services.rag.parsing.image_store import ImageStore
from app.services.rag.parsing.md_parser import extract_image_storage_keys
from app.services.rag.pipeline import _HashRagPipeline
from app.services.rag.vector_store import JsonVectorStore
from app.services.rag.vlm_service import VlmService

MINIMAL_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
    b"\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01"
    b"\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _make_docx_with_inline_picture(path: Path) -> None:
    document = docx.Document()
    document.add_heading("产品说明", level=1)
    document.add_paragraph("正文介绍键盘布局。")
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(BytesIO(MINIMAL_PNG))
    document.add_paragraph("安装步骤见下图。")
    document.save(path)


def test_docx_inline_image_in_block_text(tmp_path: Path) -> None:
    path = tmp_path / "with_image.docx"
    _make_docx_with_inline_picture(path)
    store = ImageStore(root=tmp_path)

    document = DocxDocumentParser(image_store=store).parse(path)

    assert len(document.images) == 1
    key = document.images[0].storage_key
    assert key.startswith("kb_images/")
    assert f"![图示]({key})" in document.blocks[0].text
    assert extract_image_storage_keys(document.full_text) == [key]


def test_docx_multimodal_index_with_vlm(tmp_path: Path) -> None:
    path = tmp_path / "guide.docx"
    _make_docx_with_inline_picture(path)
    upload_root = tmp_path / "uploads"
    store = ImageStore(upload_root)

    document = DocxDocumentParser(image_store=store).parse(path)
    settings = Settings(
        LOCAL_UPLOAD_ROOT=str(upload_root),
        RAG_BACKEND="local",
        LOCAL_VECTOR_STORE_PATH=str(tmp_path / "vectors"),
        VLM_ENABLED=True,
        RAG_EMBEDDING_DIMENSIONS=256,
    )
    pipeline = _HashRagPipeline(
        settings,
        JsonVectorStore(tmp_path / "vectors"),
        vlm_service=VlmService(
            settings,
            chat_completion=lambda messages: "键位图：含 Enter 与方向键",
        ),
    )
    counts = pipeline.index_data(
        knowledge_base="kb_docx",
        file_id="file-docx",
        document=document,
        config=ChunkingConfig(
            strategy="default",
            max_chunk_size=400,
            overlap=0,
            index_size=80,
            meta_headings=True,
        ),
        file_name="guide.docx",
        force_image_description=True,
    )
    assert counts["images"] >= 1
    hits = pipeline.search_data(
        knowledge_base="kb_docx",
        query="键位 Enter",
        top_k=5,
    )
    assert any(h.citation.get("type") == "image" for h in hits)
